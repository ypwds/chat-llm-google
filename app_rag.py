import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

from langchain_classic.chains import RetrievalQA
#from langchain_classic.chains import create_retrieval_chain

from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_community.embeddings import HuggingFaceEmbeddings
import PyPDF2
import docx
import io
import tempfile
import os

# Configuração da página
st.set_page_config(
    page_title="Chatbot com RAG - LangChain",
    page_icon="📚",
    layout="wide"
)

# Título da aplicação
st.title("Chatbot com RAG - LangChain")
st.markdown("*Faça upload de documentos e converse sobre o conteúdo deles*")

# Sidebar para configurações
with st.sidebar:
    st.header("⚙️ Configurações")
    
    # Campo para API Key
    api_key = st.text_input(
        "Google API Key",
        type="password",
        help="Cole aqui sua chave de API do Google AI Studio"
    )
    
    st.markdown("---")
    st.header("Parâmetros RAG")
    
    chunk_size = st.slider("Tamanho do Chunk", 500, 2000, 1000, 100)
    chunk_overlap = st.slider("Sobreposição", 50, 300, 100, 50)
    k_documents = st.slider("Documentos Recuperados", 2, 10, 4, 1)
    
    if st.button("Limpar Base de Conhecimento"):
        if 'vectorstore' in st.session_state:
            del st.session_state.vectorstore
        if 'messages' in st.session_state:
            st.session_state.messages = []
        st.success("Base de conhecimento limpa!")
        st.rerun()
    
    st.markdown("---")
    st.markdown("### 📝 Como usar:")
    st.markdown("""
    1. Insira sua API Key do Google
    2. Faça upload do(s) documento(s)
    3. Aguarde o processamento
    4. Faça perguntas sobre o conteúdo
    """)
    
    st.markdown("---")
    st.markdown("### 🔗 Links úteis:")
    st.markdown("[Obter API Key](https://aistudio.google.com/app/apikey)")

# Funções para processar diferentes tipos de arquivo
def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Erro ao processar PDF: {str(e)}")
        return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Erro ao processar DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file):
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Erro ao processar TXT: {str(e)}")
        return ""

# Função para processar documentos
def process_documents(files, api_key, chunk_size, chunk_overlap):
    if not files or not api_key:
        return None
    
    all_texts = []
    
    with st.spinner("Processando documentos..."):
        for file in files:
            file_extension = file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                text = extract_text_from_pdf(file)
            elif file_extension == 'docx':
                text = extract_text_from_docx(file)
            elif file_extension == 'txt':
                text = extract_text_from_txt(file)
            else:
                continue
            
            if text:
                # Criar documento com metadados
                doc = Document(
                    page_content=text,
                    metadata={"source": file.name, "type": file_extension}
                )
                all_texts.append(doc)
    
    if not all_texts:
        st.error("Nenhum texto foi extraído dos documentos.")
        return None
    
    # Dividir textos em chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    
    chunks = text_splitter.split_documents(all_texts)
    
    # Criar embeddings
    try:
        embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/distiluse-base-multilingual-cased-v2"
            )
        
        # Criar vectorstore
        vectorstore = FAISS.from_documents(chunks, embeddings)
        
        st.success(f"✅ {len(chunks)} chunks criados de {len(files)} arquivo(s)!")
        return vectorstore
        
    except Exception as e:
        st.error(f"Erro ao criar embeddings: {str(e)}")
        return None

# Inicializar o histórico de mensagens
if "messages" not in st.session_state:
    st.session_state.messages = []

# Verificar se a API key foi fornecida
if not api_key:
    st.warning("⚠️ Por favor, insira sua API Key do Google na barra lateral para começar.")
    st.stop()

# Seção de upload de documentos na página principal
st.header("📄 Upload de Documentos")

col_upload1, col_upload2 = st.columns([3, 1])

with col_upload1:
    uploaded_files = st.file_uploader(
        "Selecione os arquivos para envio",
        type=['pdf', 'txt', 'docx'],
        accept_multiple_files=True,
        help="Formatos suportados: PDF, TXT, DOCX"
    )

with col_upload2:
    if uploaded_files:
        st.metric("Arquivos", len(uploaded_files))
        if 'vectorstore' in st.session_state:
            st.success("✅ Processado")
        else:
            st.info("⏳ Aguardando...")

# Mostrar arquivos carregados
if uploaded_files:
    st.write("**Arquivos selecionados:**")
    cols = st.columns(min(len(uploaded_files), 4))
    for i, file in enumerate(uploaded_files):
        with cols[i % 4]:
            st.write(f"📄 {file.name}")
            st.caption(f"{file.size / 1024:.1f} KB")

st.markdown("---")

# Processar documentos se foram enviados
if uploaded_files and 'vectorstore' not in st.session_state:
    vectorstore = process_documents(uploaded_files, api_key, chunk_size, chunk_overlap)
    if vectorstore:
        st.session_state.vectorstore = vectorstore

st.header("💬 Converse sobre o conteúdo do(s) Documento(s)")

# Container para o chat
chat_container = st.container()

# Exibir mensagens do histórico
with chat_container:
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message("user"):
                st.write(message["content"])
        else:
            with st.chat_message("assistant"):
                st.write(message["content"])
                if "sources" in message:
                    with st.expander("📄 Fontes utilizadas"):
                        for source in message["sources"]:
                            st.write(f"• **{source['source']}** (Relevância: {source['score']:.2f})")
                            st.write(f"  _{source['content'][:200]}..._")

# Prompt 
custom_template = """Use o contexto fornecido para responder à pergunta, de forma educada. Se você não conseguir encontrar a resposta no contexto, diga que não sabe.

Contexto:
{context}

Pergunta: {question}

Resposta detalhada:"""

PROMPT = PromptTemplate(
    template=custom_template,
    input_variables=["context", "question"]
)

# Input do usuário
if prompt := st.chat_input("Faça uma pergunta sobre os documentos..."):
    if 'vectorstore' not in st.session_state:
        st.warning("⚠️ Por favor, faça upload de documentos primeiro!")
        st.stop()
    
    # Adicionar mensagem do usuário ao histórico
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Exibir mensagem do usuário
    with st.chat_message("user"):
        st.write(prompt)
    
    # Gerar resposta do assistente
    with st.chat_message("assistant"):
        with st.spinner("Buscando informações nos documentos..."):
            try:
                # Inicializar o modelo
                llm = ChatGoogleGenerativeAI(
                    model="gemini-2.5-flash-lite",
                    google_api_key=api_key,
                    temperature=0.3,
                    max_tokens=250
                )
                
                # Criar chain RAG
                qa_chain = RetrievalQA.from_chain_type(
                    llm=llm,
                    chain_type="stuff",
                    retriever=st.session_state.vectorstore.as_retriever(
                        search_kwargs={"k": k_documents}
                    ),
                    chain_type_kwargs={"prompt": PROMPT},
                    return_source_documents=True
                )
                
                # Executar query
                result = qa_chain({"query": prompt})
                
                # Exibir resposta
                response = result["result"]
                st.write(response)
                
                # Preparar informações das fontes
                sources = []
                for i, doc in enumerate(result["source_documents"]):
                    sources.append({
                        "source": doc.metadata.get("source", f"Documento {i+1}"),
                        "content": doc.page_content,
                        "score": 1.0 - (i * 0.1)  # Score simulado baseado na ordem
                    })
                
                # Adicionar resposta ao histórico
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": response,
                    "sources": sources
                })
                
            except Exception as e:
                st.error(f"Erro ao gerar resposta: {str(e)}")